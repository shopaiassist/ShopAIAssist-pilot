"""Handles document uploads, processing, and storage.

This module provides functionality for:
    - Validating document types (.docx and .xml).
    - Extracting text from documents.
    - Chunking text into manageable pieces.
    - Generating embeddings using OpenAI.
    - Storing documents and embeddings in S3.
    - Deleting documents from S3.

Classes:
    DocumentHandler: Handles document processing and storage operations.

Dependencies:
    - boto3: For S3 operations.
    - docx: For parsing .docx files.
    - lxml: For parsing XML files.
    - pandas: For DataFrame operations.
    - uuid: For generating unique IDs.
    - src.openai_utils: For generating embeddings.
    - src.utils: For loading configuration.
"""
import os
import uuid
import boto3
import pandas as pd
from io import BytesIO
import docx
import tiktoken
from typing import List, Dict, Tuple, BinaryIO
from fastapi import UploadFile, HTTPException
from loguru import logger
import tempfile
from pathlib import Path
from tqdm import tqdm
from langchain_text_splitters import CharacterTextSplitter
import xml.etree.ElementTree as ET
from PyPDF2 import PdfReader
from concurrent.futures import ThreadPoolExecutor

from src.utils import load_config
from src.openai_utils import OpenaiUtils
from src.skills_backend.document_qa.trlabs_smartPDFParser.file_reader import DocTreePDFReader


class FileParseError(Exception):
    """Custom exception for file parsing errors."""

    pass


class DocumentHandler:
    """Handler for document processing and storage operations."""

    def __init__(self):
        """Initialize the DocumentHandler with configuration settings."""
        self.config = load_config()
        self.document_upload_dir = self.config["document_qa"]["document_upload_dir"]
        self.s3_client = boto3.client("s3")
        self.openai_utils = OpenaiUtils()
        self.pdf_parser = DocTreePDFReader(self.config["document_qa"]["pdf_parser_url"])
        self.tokenizer = tiktoken.encoding_for_model(self.config["document_qa"]["embedding_model"])
        self.bucket_name = self.document_upload_dir.split("/")[2]
        self.max_file_size = 10 * 1024 * 1024  # 10MB in bytes
        self.supported_file_types = [".pdf", ".xml"]  # , ".docx"]

    def _generate_embedding_for_chunk(self, chunk_content: str) -> List[float]:
        """Generate embedding for a single chunk - thread-safe helper method.

        Args:
            chunk_content (str): The text content to generate embedding for.

        Returns:
            List[float]: The embedding vector.
        """
        return self.openai_utils.get_embedding(chunk_content)

    async def validate_document(self, file: UploadFile) -> bool:
        """Validate the document file type, size, and content.

        Args:
            file (UploadFile): The uploaded file object.

        Returns:
            bool: True if the document is valid, False otherwise.

        Raises:
            HTTPException: If the document is invalid.
        """
        # Check file size
        file.file.seek(0, os.SEEK_END)
        file_size = file.file.tell()
        file.file.seek(0)

        if file_size > self.max_file_size:
            raise HTTPException(
                status_code=400,
                detail=f"File size exceeds the maximum allowed size of {self.max_file_size / (1024 * 1024)}MB",
            )

        # Check file type
        file_extension = os.path.splitext(str(file.filename))[1].lower()
        if file_extension not in self.supported_file_types:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type. Supported types are: {', '.join(self.supported_file_types)}",
            )

        # Perform content-specific validation
        if file_extension == ".pdf":
            is_valid, error_message = await self.validate_pdf(file)
            if not is_valid:
                raise FileParseError(f"Invalid PDF file: {error_message}")
        elif file_extension == ".xml":
            is_valid, error_message = await self.validate_xml(file)
            if not is_valid:
                raise FileParseError(f"Invalid XML file: {error_message}")

        file.file.seek(0)
        return True

    async def validate_pdf(self, file: UploadFile) -> Tuple[bool, str]:
        """Validate PDF structure and basic readability.

        This method checks if:
        1. The PDF has a valid structure
        2. The PDF contains at least one page
        3. The PDF is not corrupted

        Args:
            file (UploadFile): The PDF file object.

        Returns:
            Tuple[bool, str]: A tuple containing (is_valid, error_message).
        """
        try:
            # Use PyPDF2 to validate the PDF structure
            reader = PdfReader(file.file, strict=True)

            # Check if the document has pages
            if len(reader.pages) == 0:
                return False, "PDF document contains no pages"

            return True, ""

        except Exception as e:
            logger.error(f"Error validating PDF: {e}")
            return False, f"PDF validation failed: {str(e)}"

    async def validate_xml(self, file: UploadFile) -> Tuple[bool, str]:
        """Validate XML structure and content.

        This method checks if the XML file is well-formed and contains actual content.

        Args:
            file_content (bytes): The XML file content.

        Returns:
            Tuple[bool, str]: A tuple containing (is_valid, error_message).
        """
        try:
            # Parse XML to check structure
            file_content = await file.read()
            root = ET.fromstring(file_content)

            # Check if XML has content
            if len(root) == 0 and not root.text:
                return False, "XML document appears to be empty or contains no meaningful content"

            return True, ""

        except ET.ParseError as e:
            return False, f"XML syntax error: {str(e)}"
        except UnicodeDecodeError as e:
            return False, f"XML file contains invalid character encoding: {str(e)}"
        except Exception as e:
            logger.error(f"Error validating XML: {e}")
            return False, f"XML validation failed: {str(e)}"

    async def pdf_embeddings(self, file: UploadFile, file_extension: str) -> pd.DataFrame:
        """Process PDF files to extract content and generate embeddings.

        This function handles PDF document processing by:
        1. Creating a temporary file from the upload
        2. Parsing the PDF document structure using the LLM Sherpa PDF parser
        3. Chunking the content based on configured chunk size
        4. Generating embeddings for each content chunk
        5. Creating a structured DataFrame with document metadata and embeddings

        Args:
            file (UploadFile): The uploaded PDF file object

        Returns:
            pd.DataFrame: DataFrame containing processed chunks with:
                - Document metadata (filename, source, content type)
                - Page information (start/end page numbers)
                - Raw text content
                - Generated embeddings
                - Unique chunk identifiers

        Raises:
            HTTPException: If PDF parsing fails or encounters errors
        """
        # Create temp directory in project root
        project_root = Path(__file__).parent.parent.parent.parent
        temp_dir = project_root / "temp"
        os.makedirs(temp_dir, exist_ok=True)

        try:
            # Create a temporary directory that will be automatically cleaned up
            with tempfile.TemporaryDirectory(dir=temp_dir) as temp_dir_path:
                temp_file_path = os.path.join(temp_dir_path, f"{file.filename}")

                with open(temp_file_path, "wb") as temp_file:
                    temp_file.write(await file.read())

                doc_tree = self.pdf_parser.get_doc_tree(
                    path_or_url=temp_file_path, file_name=file.filename, tokenizer=self.tokenizer
                )

                chunks = doc_tree.get_chunks(chunk_size=self.config["document_qa"]["chunk_size"])

                # Extract chunk contents for parallel processing
                chunk_contents = [chunk["content"] for chunk in chunks]

                # Generate embeddings in parallel while preserving order
                with ThreadPoolExecutor(max_workers=8) as executor:
                    embeddings = list(
                        tqdm(
                            executor.map(self._generate_embedding_for_chunk, chunk_contents),
                            total=len(chunk_contents),
                            desc="Generating embeddings",
                        )
                    )

                records = []
                for chunk, embedding in zip(chunks, embeddings):
                    # chunk_title = f"{file.filename} - Page {chunk['start_page']} - {chunk['end_page']}"
                    records.append(
                        {
                            "Product": "user_upload",
                            "Content": "",
                            "root_link": "",
                            "url": "",
                            "Source": "user_upload",
                            "clean_path": "",
                            "content_type": "pdf",
                            "published_date": "",
                            "file_name": file.filename,
                            "raw_text": chunk["content"],
                            "title": "",
                            "description": "",
                            "chunk_id": str(uuid.uuid4()),
                            "open_ai_embeddings": embedding,
                        }
                    )

                if len(records) == 0 or (
                    len(records) == 1 and records[0]["raw_text"].strip() == str(file.filename).strip(file_extension)
                ):
                    raise FileParseError("No valid content found in PDF file. Check the file format.")

                return pd.DataFrame(records)

        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            raise HTTPException(status_code=400, detail=f"Error parsing PDF: {str(e)}")

    async def extract_text_from_xml(self, file: UploadFile) -> pd.DataFrame:
        """Extract text from an XML file.

        Args:
            file (UploadFile): The FastAPI file object.

        Returns:
            pd.DataFrame: A DataFrame containing the extracted text and embeddings.
        """
        try:
            splitter = CharacterTextSplitter.from_tiktoken_encoder(
                model_name=self.config["document_qa"]["embedding_model"],
                chunk_size=self.config["document_qa"]["chunk_size"],
                chunk_overlap=0,
                separator="\n",
            )
            xml_text = await file.read()
            content_chunks = splitter.split_text(xml_text.decode("utf-8"))

            # Generate embeddings in parallel while preserving order
            with ThreadPoolExecutor(max_workers=8) as executor:
                embeddings = list(
                    tqdm(
                        executor.map(self._generate_embedding_for_chunk, content_chunks),
                        total=len(content_chunks),
                        desc="Generating embeddings",
                    )
                )

            records = []
            for chunk, embedding in zip(content_chunks, embeddings):
                records.append(
                    {
                        "Product": "user_upload",
                        "Content": "",
                        "root_link": "",
                        "url": "",
                        "Source": "user_upload",
                        "clean_path": "",
                        "content_type": "pdf",
                        "published_date": "",
                        "file_name": file.filename,
                        "raw_text": chunk,
                        "title": "",
                        "description": file.filename,
                        "chunk_id": str(uuid.uuid4()),
                        "open_ai_embeddings": embedding,
                    }
                )

            if len(records) == 0:
                raise FileParseError("No valid content found in XML file. Check the file format.")

            return pd.DataFrame(records)

        except Exception as e:
            logger.error(f"Error extracting text from XML: {e}")
            raise HTTPException(status_code=400, detail=f"Invalid XML file: {str(e)}")

    async def save_to_s3(
        self,
        dataframe: pd.DataFrame,
        tenant_id: str,
        user_id: str,
        document_id: str,
        filename: str | None,
    ) -> str:
        """Save the original document and embeddings DataFrame to S3.

        Args:
            file_content (BinaryIO): The original file content.
            dataframe (pd.DataFrame): The DataFrame with chunks and embeddings.
            tenant_id (str): The tenant ID.
            user_id (str): The user ID.
            document_id (str): The document ID.
            filename (str): The original filename.

        Returns:
            Dict[str, str]: A dictionary with the S3 paths of the saved files.
        """
        # Construct S3 paths
        base_path = f"tenant_id={tenant_id}/user_id={user_id}"
        pickle_key = f"{base_path}/{document_id}_{os.path.splitext(str(filename))[0]}.pkl"

        # Save DataFrame as pickle
        pickle_buffer = BytesIO()
        dataframe.to_pickle(pickle_buffer)
        pickle_buffer.seek(0)
        self.s3_client.upload_fileobj(pickle_buffer, self.bucket_name, pickle_key)

        return f"s3://{self.bucket_name}/{pickle_key}"

    async def delete_from_s3(self, tenant_id: str, user_id: str, document_id: str) -> bool:
        """Delete document files from S3.

        Args:
            tenant_id (str): The tenant ID.
            user_id (str): The user ID.
            document_id (str): The document ID.

        Returns:
            bool: True if deletion was successful, False otherwise.
        """
        try:
            # List objects with the document_id prefix
            base_path = f"tenant_id={tenant_id}/user_id={user_id}"
            response = self.s3_client.list_objects_v2(Bucket=self.bucket_name, Prefix=f"{base_path}/{document_id}")

            # Check if any objects were found
            if "Contents" not in response:
                return False

            # Delete each object
            for obj in response["Contents"]:
                self.s3_client.delete_object(Bucket=self.bucket_name, Key=obj["Key"])

            return True
        except Exception as e:
            logger.error(f"Error deleting document from S3: {e}")
            return False

    async def process_document(self, file: UploadFile, tenant_id: str, user_id: str) -> Dict[str, str | None]:
        """Process a document file, generate embeddings, and save to S3.

        Args:
            file (UploadFile): The uploaded file.
            tenant_id (str): The tenant ID.
            user_id (str): The user ID.

        Returns:
            Dict[str, str]: A response with document ID and status.
        """
        try:
            # Validate document
            await self.validate_document(file)

            # Generate a unique document ID
            document_id = str(uuid.uuid4())

            # Get file extension
            file_extension = os.path.splitext(str(file.filename))[1].lower()

            # # Extract text based on file type
            # if file_extension == ".docx":
            #     full_text, title, description = await self.extract_text_from_docx(BytesIO(await file.read()))
            #     # Chunk the text
            #     # Generate embeddings
            #     embeddings = await self.generate_embeddings(chunks)

            #     # Create DataFrame
            #     df = await self.create_dataframe(
            #         chunks,
            #         embeddings,
            #         str(file.filename),
            #         file_extension[1:],  # Remove the dot from extension
            #         title,
            #         description,
            #     )

            if file_extension == ".xml":
                df = await self.extract_text_from_xml(file)

            elif file_extension == ".pdf":
                df = await self.pdf_embeddings(file, file_extension)

            else:
                raise HTTPException(status_code=400, detail="Unsupported file type")

            # Save to S3
            s3_path = await self.save_to_s3(df, tenant_id, user_id, document_id, file.filename)

            return {
                "document_id": document_id,
                "filename": file.filename,
                "status": "success",
                "message": "Document processed and saved successfully",
                "s3_path": s3_path,
            }
        except HTTPException as e:
            raise e
        except FileParseError as e:
            logger.error(f"File parsing error: {e}")
            raise HTTPException(status_code=400, detail=str(e))
        except Exception as e:
            logger.error(f"Error processing document: {e}")
            raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")

    async def extract_text_from_docx(self, file_content: BinaryIO) -> Tuple[str, str, str]:
        """Extract text from a .docx file.

        Args:
            file_content (BinaryIO): The file content as a binary stream.

        Returns:
            Tuple[str, str, str]: A tuple containing the extracted text, title, and description.
        """
        doc = docx.Document(file_content)

        # Extract title (assume first paragraph is title if it's short)
        title = doc.paragraphs[0].text if doc.paragraphs and len(doc.paragraphs[0].text) < 100 else ""

        # Extract description (assume second paragraph is description if it's not too long)
        description = doc.paragraphs[1].text if len(doc.paragraphs) > 1 and len(doc.paragraphs[1].text) < 500 else ""

        # Extract full text
        full_text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])

        return full_text, title, description

    async def chunk_text(self, text: str, chunk_size: int = 1024, overlap: int = 100) -> List[str]:
        """Split text into chunks of specified size with overlap.

        Args:
            text (str): The text to chunk.
            chunk_size (int, optional): The maximum size of each chunk. Defaults to 1000.
            overlap (int, optional): The overlap between chunks. Defaults to 100.

        Returns:
            List[str]: A list of text chunks.
        """
        if not text:
            return []

        # Split text into paragraphs
        paragraphs = text.split("\n")

        chunks: List[str] = []
        current_chunk: List[str] = []
        current_size = 0

        for paragraph in paragraphs:
            paragraph_size = len(paragraph)

            # If adding this paragraph would exceed the chunk size, finalize the current chunk
            if current_size + paragraph_size > chunk_size and current_chunk:
                chunks.append("\n".join(current_chunk))

                # Keep some paragraphs for overlap
                overlap_size = 0
                overlap_paragraphs: List[str] = []

                for p in reversed(current_chunk):
                    if overlap_size + len(p) <= overlap:
                        overlap_paragraphs.insert(0, p)
                        overlap_size += len(p)
                    else:
                        break

                current_chunk = overlap_paragraphs
                current_size = overlap_size

            current_chunk.append(paragraph)
            current_size += paragraph_size

        # Add the last chunk if it's not empty
        if current_chunk:
            chunks.append("\n".join(current_chunk))

        return chunks

    async def generate_embeddings(self, chunks: List[str]) -> List[List[float]]:
        """Generate embeddings for text chunks using OpenAI.

        Args:
            chunks (List[str]): The text chunks to generate embeddings for.

        Returns:
            List[List[float]]: A list of embedding vectors.
        """
        embeddings = []

        for chunk in chunks:
            embedding = self.openai_utils.get_embedding(chunk)
            embeddings.append(embedding)

        return embeddings

    async def create_dataframe(
        self,
        chunks: List[str],
        embeddings: List[List[float]],
        filename: str,
        file_type: str,
        title: str = "",
        description: str = "",
    ) -> pd.DataFrame:
        """Create a pandas DataFrame with document chunks and embeddings.

        Args:
            chunks (List[str]): The text chunks.
            embeddings (List[List[float]]): The embedding vectors.
            filename (str): The original filename.
            file_type (str): The file type (e.g., "DOCX" or "XML").
            title (str, optional): The document title. Defaults to "".
            description (str, optional): The document description. Defaults to "".

        Returns:
            pd.DataFrame: A DataFrame containing the document chunks and embeddings.
        """
        data: Dict[str, list] = {
            "Product": [],
            "Content": [],
            "root_link": [],
            "url": [],
            "Source": [],
            "clean_path": [],
            "content_type": [],
            "file_name": [],
            "raw_text": [],
            "title": [],
            "description": [],
            "open_ai_embeddings": [],
            "chunk_id": [],
        }

        for chunk, embedding in zip(chunks, embeddings):
            data["Product"].append(title if title else filename)  # Use title as product if available
            data["Content"].append(chunk)
            data["root_link"].append("")
            data["url"].append("")
            data["Source"].append(file_type.upper())
            data["clean_path"].append("")
            data["content_type"].append("AEM")
            data["file_name"].append(filename)
            data["raw_text"].append(chunk)
            data["title"].append(title if title else filename)
            data["description"].append(description)
            data["open_ai_embeddings"].append(embedding)
            data["chunk_id"].append(str(uuid.uuid4()))

        return pd.DataFrame(data)
